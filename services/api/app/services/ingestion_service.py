"""
Document ingestion service for DOCX to OSCAL mapping.

Provides functionality to ingest legacy document formats (primarily DOCX)
and map them to OSCAL structures, particularly for SSP generation.
"""

import re
import json
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, Tuple
from dataclasses import dataclass
from datetime import datetime, timezone
from uuid import uuid4
import xml.etree.ElementTree as ET

import structlog
from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table

logger = structlog.get_logger()


@dataclass
class IngestionResult:
    """Result of document ingestion process."""
    success: bool
    document_type: str  # "ssp", "sap", "sar", etc.
    oscal_document: Optional[Dict[str, Any]] = None
    extracted_content: Optional[Dict[str, Any]] = None
    issues: List[str] = None
    metadata: Optional[Dict[str, Any]] = None
    processing_time_ms: Optional[int] = None
    
    def __post_init__(self):
        if self.issues is None:
            self.issues = []


@dataclass
class ControlMapping:
    """Mapping of extracted control information."""
    control_id: str
    control_title: str
    implementation_statement: str
    responsible_role: Optional[str] = None
    implementation_status: str = "implemented"
    control_origination: List[str] = None
    
    def __post_init__(self):
        if self.control_origination is None:
            self.control_origination = ["system-specific"]


class DocumentStructureAnalyzer:
    """
    Analyzes document structure to identify content patterns.
    
    Uses various heuristics to identify sections, controls, tables,
    and other structured content in DOCX documents.
    """
    
    def __init__(self):
        self.logger = structlog.get_logger().bind(component="doc_analyzer")
        
        # Common patterns for identifying different content types
        self.control_patterns = [
            re.compile(r'^([A-Z]{2}-\d+(?:\(\d+\))?)\s*[-–—]\s*(.+?)(?:\s*\(.*\))?$', re.IGNORECASE),
            re.compile(r'^([A-Z]{2}-\d+(?:\.\d+)?)\s+(.+)', re.IGNORECASE),
            re.compile(r'Control\s+([A-Z]{2}-\d+)', re.IGNORECASE),
        ]
        
        self.section_patterns = [
            re.compile(r'^\d+\.\s+(.+)', re.IGNORECASE),  # Numbered sections
            re.compile(r'^[A-Z]\.\s+(.+)', re.IGNORECASE),  # Letter sections
            re.compile(r'^(System\s+Characteristics|Control\s+Implementation|Authorization\s+Boundary)', re.IGNORECASE),
        ]
        
        self.role_patterns = [
            re.compile(r'Responsible\s+Role[s]?\s*[:\-]\s*(.+)', re.IGNORECASE),
            re.compile(r'Implementation\s+Role[s]?\s*[:\-]\s*(.+)', re.IGNORECASE),
            re.compile(r'Role\s*[:\-]\s*(.+)', re.IGNORECASE),
        ]
    
    def analyze_document_structure(self, doc: DocxDocument) -> Dict[str, Any]:
        """
        Analyze the overall structure of the document.
        
        Returns:
            Dictionary with structure analysis results
        """
        structure = {
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "headings": [],
            "sections": {},
            "controls_identified": [],
            "document_type": "unknown",
            "potential_ssp_sections": [],
        }
        
        # Analyze headings and section structure
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check for headings (typically have specific styles)
            if paragraph.style.name.startswith('Heading') or self._is_likely_heading(text):
                structure["headings"].append({
                    "level": self._extract_heading_level(paragraph.style.name),
                    "text": text,
                    "paragraph_index": i,
                })
        
        # Identify potential SSP sections
        structure["potential_ssp_sections"] = self._identify_ssp_sections(structure["headings"])
        
        # Identify controls in content
        structure["controls_identified"] = self._identify_controls_in_document(doc)
        
        # Determine document type based on structure
        structure["document_type"] = self._determine_document_type(structure)
        
        return structure
    
    def _is_likely_heading(self, text: str) -> bool:
        """Check if text is likely a heading based on patterns."""
        # Check for numbered sections
        if re.match(r'^\d+\.\s+[A-Z]', text):
            return True
        
        # Check for common SSP section titles
        ssp_sections = [
            'system characteristics', 'system implementation', 'control implementation',
            'authorization boundary', 'network architecture', 'data flow',
            'system security plan', 'responsible roles', 'system description'
        ]
        
        return any(section in text.lower() for section in ssp_sections)
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        if 'Heading' in style_name:
            match = re.search(r'Heading\s*(\d+)', style_name)
            if match:
                return int(match.group(1))
        return 1
    
    def _identify_ssp_sections(self, headings: List[Dict]) -> List[Dict]:
        """Identify sections that likely correspond to SSP components."""
        ssp_section_mappings = {
            'system characteristics': 'system-characteristics',
            'system description': 'system-characteristics',
            'authorization boundary': 'system-characteristics.authorization-boundary',
            'network architecture': 'system-characteristics.network-architecture',
            'data flow': 'system-characteristics.data-flow',
            'system implementation': 'system-implementation',
            'system components': 'system-implementation.components',
            'control implementation': 'control-implementation',
            'controls': 'control-implementation',
            'security controls': 'control-implementation',
            'responsible roles': 'metadata.roles',
            'responsible parties': 'metadata.responsible-parties',
        }
        
        identified_sections = []
        for heading in headings:
            heading_text = heading["text"].lower()
            for pattern, oscal_path in ssp_section_mappings.items():
                if pattern in heading_text:
                    identified_sections.append({
                        "heading": heading,
                        "section_type": pattern,
                        "oscal_path": oscal_path,
                        "confidence": self._calculate_section_confidence(heading_text, pattern)
                    })
        
        return identified_sections
    
    def _calculate_section_confidence(self, text: str, pattern: str) -> float:
        """Calculate confidence score for section identification."""
        if text == pattern:
            return 1.0
        elif pattern in text and len(text.split()) <= 5:
            return 0.8
        elif pattern in text:
            return 0.6
        else:
            return 0.4
    
    def _identify_controls_in_document(self, doc: DocxDocument) -> List[Dict]:
        """Identify security controls mentioned in the document."""
        controls = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check for control patterns
            for pattern in self.control_patterns:
                match = pattern.match(text)
                if match:
                    control_id = match.group(1).upper()
                    control_title = match.group(2).strip() if len(match.groups()) > 1 else ""
                    
                    controls.append({
                        "control_id": control_id,
                        "control_title": control_title,
                        "paragraph_index": i,
                        "paragraph_text": text,
                        "pattern_used": pattern.pattern
                    })
        
        # Remove duplicates based on control_id
        unique_controls = {}
        for control in controls:
            if control["control_id"] not in unique_controls:
                unique_controls[control["control_id"]] = control
        
        return list(unique_controls.values())
    
    def _determine_document_type(self, structure: Dict[str, Any]) -> str:
        """Determine document type based on structural analysis."""
        
        # Check for SSP indicators
        ssp_indicators = [
            "system security plan", "control implementation", "authorization boundary",
            "system characteristics", "system implementation"
        ]
        
        # Check for SAP indicators  
        sap_indicators = [
            "security assessment plan", "assessment procedures", "assessment methods"
        ]
        
        # Check for SAR indicators
        sar_indicators = [
            "security assessment report", "assessment results", "findings", "vulnerabilities"
        ]
        
        # Check for POA&M indicators
        poam_indicators = [
            "plan of action", "milestones", "remediation", "risk mitigation"
        ]
        
        headings_text = " ".join([h["text"].lower() for h in structure["headings"]])
        controls_count = len(structure["controls_identified"])
        
        # Score each document type
        scores = {
            "ssp": sum(1 for indicator in ssp_indicators if indicator in headings_text),
            "sap": sum(1 for indicator in sap_indicators if indicator in headings_text),
            "sar": sum(1 for indicator in sar_indicators if indicator in headings_text),
            "poam": sum(1 for indicator in poam_indicators if indicator in headings_text),
        }
        
        # Bonus for controls (indicates SSP)
        if controls_count > 10:
            scores["ssp"] += 2
        elif controls_count > 5:
            scores["ssp"] += 1
        
        # Return highest scoring type, or "unknown" if no clear winner
        if max(scores.values()) > 0:
            return max(scores, key=scores.get)
        else:
            return "unknown"


class DOCXToOSCALMapper:
    """
    Maps DOCX content to OSCAL document structures.
    
    Takes analyzed document content and maps it to appropriate
    OSCAL document formats, primarily SSPs.
    """
    
    def __init__(self):
        self.logger = structlog.get_logger().bind(component="docx_mapper")
        self.analyzer = DocumentStructureAnalyzer()
    
    def map_to_ssp(
        self, 
        doc: DocxDocument, 
        document_title: str = "Imported System Security Plan",
        system_id: str = None
    ) -> Dict[str, Any]:
        """
        Map DOCX content to OSCAL SSP format.
        
        Args:
            doc: DOCX document object
            document_title: Title for the SSP
            system_id: System identifier
            
        Returns:
            OSCAL SSP document structure
        """
        if not system_id:
            system_id = f"imported-system-{uuid4()}"
        
        # Analyze document structure
        structure = self.analyzer.analyze_document_structure(doc)
        
        # Extract content sections
        extracted_content = self._extract_content_sections(doc, structure)
        
        # Build OSCAL SSP
        ssp = {
            "system-security-plan": {
                "uuid": str(uuid4()),
                "metadata": self._build_metadata(document_title, extracted_content),
                "system-characteristics": self._build_system_characteristics(extracted_content, system_id),
                "system-implementation": self._build_system_implementation(extracted_content),
                "control-implementation": self._build_control_implementation(extracted_content, structure),
                "back-matter": self._build_back_matter()
            }
        }
        
        return ssp
    
    def _extract_content_sections(self, doc: DocxDocument, structure: Dict[str, Any]) -> Dict[str, Any]:
        """Extract content from different document sections."""
        content = {
            "system_description": "",
            "authorization_boundary": "",
            "network_architecture": "",
            "data_types": [],
            "system_components": [],
            "responsible_roles": [],
            "control_implementations": {},
            "tables": [],
            "raw_paragraphs": []
        }
        
        current_section = None
        section_content = []
        
        # Process paragraphs to group content by sections
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            content["raw_paragraphs"].append({
                "index": i,
                "text": text,
                "style": paragraph.style.name
            })
            
            # Check if this is a new section heading
            if self.analyzer._is_likely_heading(text):
                # Store previous section content
                if current_section and section_content:
                    content[current_section] = " ".join(section_content)
                
                # Identify new section
                current_section = self._map_heading_to_content_key(text)
                section_content = []
            else:
                # Add to current section
                section_content.append(text)
        
        # Handle last section
        if current_section and section_content:
            content[current_section] = " ".join(section_content)
        
        # Extract information from tables
        content["tables"] = self._extract_table_content(doc.tables)
        
        # Extract controls from identified controls
        for control in structure.get("controls_identified", []):
            control_id = control["control_id"]
            content["control_implementations"][control_id] = {
                "control_id": control_id,
                "title": control["control_title"],
                "implementation_statement": self._extract_control_implementation(doc, control),
                "responsible_role": self._extract_responsible_role(doc, control),
            }
        
        return content
    
    def _map_heading_to_content_key(self, heading_text: str) -> str:
        """Map heading text to content dictionary key."""
        heading_lower = heading_text.lower()
        
        mappings = {
            "system description": "system_description",
            "system characteristics": "system_description", 
            "authorization boundary": "authorization_boundary",
            "network architecture": "network_architecture",
            "system components": "system_components_text",
            "responsible roles": "responsible_roles_text",
            "data types": "data_types_text",
        }
        
        for pattern, key in mappings.items():
            if pattern in heading_lower:
                return key
        
        return "miscellaneous"
    
    def _extract_table_content(self, tables: List[Table]) -> List[Dict]:
        """Extract structured content from document tables."""
        table_data = []
        
        for table_index, table in enumerate(tables):
            table_content = {
                "index": table_index,
                "headers": [],
                "rows": []
            }
            
            for row_index, row in enumerate(table.rows):
                row_cells = [cell.text.strip() for cell in row.cells]
                
                if row_index == 0:
                    # Assume first row is headers
                    table_content["headers"] = row_cells
                else:
                    table_content["rows"].append(row_cells)
            
            # Try to identify table purpose based on headers
            table_content["purpose"] = self._identify_table_purpose(table_content["headers"])
            table_data.append(table_content)
        
        return table_data
    
    def _identify_table_purpose(self, headers: List[str]) -> str:
        """Identify the purpose of a table based on its headers."""
        headers_text = " ".join(headers).lower()
        
        if any(term in headers_text for term in ["control", "implementation", "status"]):
            return "control_implementation"
        elif any(term in headers_text for term in ["component", "system", "description"]):
            return "system_components"
        elif any(term in headers_text for term in ["role", "responsibility", "contact"]):
            return "responsible_parties"
        elif any(term in headers_text for term in ["data", "type", "classification"]):
            return "data_types"
        else:
            return "general"
    
    def _extract_control_implementation(self, doc: DocxDocument, control_info: Dict) -> str:
        """Extract implementation statement for a specific control."""
        start_index = control_info["paragraph_index"]
        implementation_text = []
        
        # Look at paragraphs following the control heading
        for i in range(start_index + 1, min(start_index + 10, len(doc.paragraphs))):
            if i < len(doc.paragraphs):
                paragraph_text = doc.paragraphs[i].text.strip()
                
                # Stop if we hit another control or major heading
                if (any(pattern.match(paragraph_text) for pattern in self.analyzer.control_patterns) or
                    self.analyzer._is_likely_heading(paragraph_text)):
                    break
                
                if paragraph_text:
                    implementation_text.append(paragraph_text)
        
        return " ".join(implementation_text) if implementation_text else f"Implementation details for {control_info['control_id']} not found in document."
    
    def _extract_responsible_role(self, doc: DocxDocument, control_info: Dict) -> Optional[str]:
        """Extract responsible role for a control."""
        start_index = control_info["paragraph_index"]
        
        # Look for role information near the control
        for i in range(max(0, start_index - 2), min(start_index + 5, len(doc.paragraphs))):
            if i < len(doc.paragraphs):
                paragraph_text = doc.paragraphs[i].text.strip()
                
                for pattern in self.analyzer.role_patterns:
                    match = pattern.search(paragraph_text)
                    if match:
                        return match.group(1).strip()
        
        return None
    
    def _build_metadata(self, title: str, content: Dict[str, Any]) -> Dict[str, Any]:
        """Build OSCAL metadata section."""
        return {
            "title": title,
            "last-modified": datetime.now(timezone.utc).isoformat(),
            "version": "1.0",
            "oscal-version": "1.1.3",
            "roles": [
                {
                    "id": "system-owner",
                    "title": "System Owner"
                },
                {
                    "id": "authorizing-official",
                    "title": "Authorizing Official"
                }
            ],
            "parties": [
                {
                    "uuid": str(uuid4()),
                    "type": "organization",
                    "name": "Organization Name (extracted from document)"
                }
            ]
        }
    
    def _build_system_characteristics(self, content: Dict[str, Any], system_id: str) -> Dict[str, Any]:
        """Build OSCAL system-characteristics section."""
        return {
            "system-id": system_id,
            "system-name": "Imported System",
            "description": content.get("system_description", "System description extracted from imported document."),
            "authorization-boundary": {
                "description": content.get("authorization_boundary", "Authorization boundary description not found in source document.")
            },
            "network-architecture": {
                "description": content.get("network_architecture", "Network architecture description not found in source document.")
            } if content.get("network_architecture") else None,
            "data-flow": {
                "description": "Data flow information extracted from source document."
            } if content.get("data_types_text") else None
        }
    
    def _build_system_implementation(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Build OSCAL system-implementation section."""
        components = []
        
        # Extract components from tables if available
        for table in content.get("tables", []):
            if table["purpose"] == "system_components":
                for row in table["rows"]:
                    if len(row) >= 2:  # At least name and description
                        components.append({
                            "uuid": str(uuid4()),
                            "type": "system",  # Default type
                            "title": row[0],
                            "description": row[1] if len(row) > 1 else row[0]
                        })
        
        # If no components found in tables, create a generic one
        if not components:
            components.append({
                "uuid": str(uuid4()),
                "type": "system",
                "title": "Imported System Component",
                "description": "System component information extracted from imported document."
            })
        
        return {
            "components": components
        }
    
    def _build_control_implementation(self, content: Dict[str, Any], structure: Dict[str, Any]) -> Dict[str, Any]:
        """Build OSCAL control-implementation section."""
        implemented_requirements = []
        
        for control_id, control_info in content.get("control_implementations", {}).items():
            implemented_requirements.append({
                "uuid": str(uuid4()),
                "control-id": control_id.lower(),
                "statements": [
                    {
                        "statement-id": f"{control_id.lower()}_stmt",
                        "uuid": str(uuid4()),
                        "description": control_info["implementation_statement"]
                    }
                ],
                "responsible-roles": [
                    {
                        "role-id": "system-owner"  # Default role
                    }
                ] if not control_info.get("responsible_role") else [
                    {
                        "role-id": "system-owner",
                        "party-uuids": []  # Would be populated with actual party UUIDs
                    }
                ]
            })
        
        return {
            "description": "Control implementations extracted from imported document.",
            "implemented-requirements": implemented_requirements
        }
    
    def _build_back_matter(self) -> Dict[str, Any]:
        """Build OSCAL back-matter section."""
        return {
            "resources": [
                {
                    "uuid": str(uuid4()),
                    "title": "Original Document",
                    "description": "Source document from which this OSCAL SSP was generated",
                    "props": [
                        {
                            "name": "document-type",
                            "value": "source-document"
                        }
                    ]
                }
            ]
        }


class DocumentIngestionService:
    """
    Main document ingestion service.
    
    Orchestrates the ingestion process from DOCX files to OSCAL documents.
    """
    
    def __init__(self):
        self.logger = structlog.get_logger().bind(component="ingestion_service")
        self.mapper = DOCXToOSCALMapper()
    
    async def ingest_docx(
        self, 
        file_path: Union[str, Path],
        target_document_type: str = "ssp",
        system_id: Optional[str] = None,
        document_title: Optional[str] = None
    ) -> IngestionResult:
        """
        Ingest a DOCX file and convert it to OSCAL format.
        
        Args:
            file_path: Path to DOCX file
            target_document_type: Target OSCAL document type
            system_id: System identifier for SSPs
            document_title: Document title override
            
        Returns:
            Ingestion result with OSCAL document
        """
        file_path = Path(file_path)
        start_time = datetime.now(timezone.utc)
        
        self.logger.info(
            "Starting document ingestion",
            file_path=str(file_path),
            target_type=target_document_type
        )
        
        try:
            # Validate file exists and is DOCX
            if not file_path.exists():
                return IngestionResult(
                    success=False,
                    document_type=target_document_type,
                    issues=[f"File not found: {file_path}"]
                )
            
            if file_path.suffix.lower() not in ['.docx']:
                return IngestionResult(
                    success=False,
                    document_type=target_document_type,
                    issues=[f"Unsupported file format: {file_path.suffix}"]
                )
            
            # Load DOCX document
            try:
                doc = Document(file_path)
            except Exception as e:
                return IngestionResult(
                    success=False,
                    document_type=target_document_type,
                    issues=[f"Failed to load DOCX file: {str(e)}"]
                )
            
            # Set document title
            if not document_title:
                document_title = file_path.stem.replace('_', ' ').replace('-', ' ').title()
            
            # Convert based on target document type
            if target_document_type == "ssp":
                oscal_doc = self.mapper.map_to_ssp(doc, document_title, system_id)
            else:
                return IngestionResult(
                    success=False,
                    document_type=target_document_type,
                    issues=[f"Document type '{target_document_type}' not yet supported"]
                )
            
            # Extract metadata about the ingestion
            structure = self.mapper.analyzer.analyze_document_structure(doc)
            
            duration = datetime.now(timezone.utc) - start_time
            
            return IngestionResult(
                success=True,
                document_type=target_document_type,
                oscal_document=oscal_doc,
                extracted_content=structure,
                metadata={
                    "source_file": str(file_path),
                    "source_file_size": file_path.stat().st_size,
                    "document_title": document_title,
                    "system_id": system_id,
                    "paragraphs_processed": structure.get("total_paragraphs", 0),
                    "tables_processed": structure.get("total_tables", 0),
                    "controls_identified": len(structure.get("controls_identified", [])),
                    "sections_identified": len(structure.get("potential_ssp_sections", [])),
                    "detected_document_type": structure.get("document_type", "unknown"),
                    "ingestion_date": start_time.isoformat(),
                },
                processing_time_ms=int(duration.total_seconds() * 1000)
            )
            
        except Exception as e:
            self.logger.error("Document ingestion failed", file_path=str(file_path), error=str(e))
            return IngestionResult(
                success=False,
                document_type=target_document_type,
                issues=[f"Ingestion failed: {str(e)}"],
                processing_time_ms=int((datetime.now(timezone.utc) - start_time).total_seconds() * 1000)
            )
    
    async def validate_ingested_document(
        self, 
        oscal_document: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Validate an ingested OSCAL document.
        
        Args:
            oscal_document: OSCAL document to validate
            
        Returns:
            Validation results
        """
        issues = []
        
        # Basic structure validation
        if "system-security-plan" in oscal_document:
            ssp = oscal_document["system-security-plan"]
            
            required_sections = ["uuid", "metadata", "system-characteristics", "control-implementation"]
            for section in required_sections:
                if section not in ssp:
                    issues.append(f"Missing required section: {section}")
            
            # Validate UUID format
            if "uuid" in ssp:
                try:
                    from uuid import UUID
                    UUID(ssp["uuid"])
                except ValueError:
                    issues.append("Invalid UUID format")
            
            # Check for control implementations
            control_impl = ssp.get("control-implementation", {})
            implemented_reqs = control_impl.get("implemented-requirements", [])
            if not implemented_reqs:
                issues.append("No control implementations found")
            
        return {
            "is_valid": len(issues) == 0,
            "issues": issues,
            "total_issues": len(issues)
        }