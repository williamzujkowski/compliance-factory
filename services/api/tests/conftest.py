"""
Test configuration and fixtures for the OSCAL Compliance Factory API.

Provides common test fixtures, database setup, and test utilities.
"""

import asyncio
import json
import tempfile
from pathlib import Path
from typing import AsyncGenerator, Dict, Any, Generator
from uuid import uuid4

import pytest
import pytest_asyncio
from fastapi.testclient import TestClient
from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine, async_sessionmaker
from sqlalchemy.pool import StaticPool

from app.main import app
from app.models import Base
from app.core.database import get_db_session
from app.core.config import get_settings


# Test database URL (in-memory SQLite for speed)
TEST_DATABASE_URL = "sqlite+aiosqlite:///:memory:"


@pytest.fixture(scope="session")
def event_loop():
    """Create an instance of the default event loop for the test session."""
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()


@pytest_asyncio.fixture
async def test_engine():
    """Create test database engine."""
    engine = create_async_engine(
        TEST_DATABASE_URL,
        echo=False,
        poolclass=StaticPool,
        connect_args={
            "check_same_thread": False,
        },
    )
    
    # Create all tables
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    
    yield engine
    
    # Clean up
    await engine.dispose()


@pytest_asyncio.fixture
async def test_session(test_engine) -> AsyncGenerator[AsyncSession, None]:
    """Create test database session."""
    async_session = async_sessionmaker(
        test_engine,
        class_=AsyncSession,
        expire_on_commit=False
    )
    
    async with async_session() as session:
        yield session


@pytest.fixture
def test_client(test_session) -> Generator[TestClient, None, None]:
    """Create test client with test database."""
    def get_test_db():
        return test_session
    
    app.dependency_overrides[get_db_session] = get_test_db
    
    with TestClient(app) as client:
        yield client
    
    app.dependency_overrides.clear()


@pytest.fixture
def sample_ssp() -> Dict[str, Any]:
    """Sample OSCAL SSP document for testing."""
    return {
        "system-security-plan": {
            "uuid": str(uuid4()),
            "metadata": {
                "title": "Test System Security Plan",
                "last-modified": "2024-01-15T10:30:00Z",
                "version": "1.0",
                "oscal-version": "1.1.3",
                "roles": [
                    {
                        "id": "system-owner",
                        "title": "System Owner"
                    },
                    {
                        "id": "isso",
                        "title": "Information System Security Officer"
                    }
                ],
                "parties": [
                    {
                        "uuid": str(uuid4()),
                        "type": "organization",
                        "name": "Test Organization"
                    }
                ]
            },
            "system-characteristics": {
                "system-id": "test-system-001",
                "system-name": "Test Information System",
                "description": "Test system for compliance factory testing.",
                "authorization-boundary": {
                    "description": "Test authorization boundary description."
                }
            },
            "system-implementation": {
                "components": [
                    {
                        "uuid": str(uuid4()),
                        "type": "software",
                        "title": "Test Application Server",
                        "description": "Primary application server for testing",
                        "status": {"state": "operational"}
                    }
                ]
            },
            "control-implementation": {
                "implemented-requirements": [
                    {
                        "uuid": str(uuid4()),
                        "control-id": "ac-2",
                        "statements": [
                            {
                                "statement-id": "ac-2_stmt",
                                "uuid": str(uuid4()),
                                "description": "Test implementation for AC-2 Account Management"
                            }
                        ]
                    },
                    {
                        "uuid": str(uuid4()),
                        "control-id": "ac-3",
                        "statements": [
                            {
                                "statement-id": "ac-3_stmt", 
                                "uuid": str(uuid4()),
                                "description": "Test implementation for AC-3 Access Enforcement"
                            }
                        ]
                    }
                ]
            }
        }
    }


@pytest.fixture
def sample_invalid_ssp() -> Dict[str, Any]:
    """Invalid OSCAL SSP document for error testing."""
    return {
        "system-security-plan": {
            # Missing required uuid
            "metadata": {
                "title": "Invalid SSP",
                # Missing required last-modified
                "version": "1.0"
                # Missing oscal-version
            },
            # Missing system-characteristics
            # Missing control-implementation
        }
    }


@pytest.fixture
def temp_oscal_file(sample_ssp) -> Generator[Path, None, None]:
    """Create temporary OSCAL file for testing."""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as f:
        json.dump(sample_ssp, f, indent=2)
        temp_path = Path(f.name)
    
    yield temp_path
    
    # Clean up
    temp_path.unlink(missing_ok=True)


@pytest.fixture
def temp_docx_file() -> Generator[Path, None, None]:
    """Create temporary DOCX file for testing."""
    from docx import Document
    
    # Create a simple test DOCX
    doc = Document()
    doc.add_heading('System Security Plan', 0)
    doc.add_heading('System Description', level=1)
    doc.add_paragraph('This is a test system for compliance validation.')
    
    doc.add_heading('Control Implementation', level=1)
    doc.add_heading('AC-2 - Account Management', level=2)
    doc.add_paragraph('The organization manages information system accounts including establishment, activation, modification, review, and removal of accounts.')
    
    doc.add_heading('AC-3 - Access Enforcement', level=2)
    doc.add_paragraph('The information system enforces approved authorizations for logical access.')
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        temp_path = Path(f.name)
    
    yield temp_path
    
    # Clean up
    temp_path.unlink(missing_ok=True)


@pytest.fixture
def mock_storage_service():
    """Mock storage service for testing."""
    from unittest.mock import Mock, AsyncMock
    from app.services.storage_service import StorageResult
    
    mock_service = Mock()
    mock_service.store_artifact = AsyncMock(return_value=StorageResult(
        success=True,
        bucket="test-bucket",
        object_key="test/object/key",
        url="https://test-bucket.s3.amazonaws.com/test/object/key",
        checksum="abc123def456",
        file_size_bytes=1024
    ))
    mock_service.get_download_url = AsyncMock(return_value="https://test-download-url.com")
    mock_service.health_check = AsyncMock(return_value={"status": "healthy"})
    
    return mock_service


@pytest.fixture
def mock_oscal_service():
    """Mock OSCAL service for testing.""" 
    from unittest.mock import Mock, AsyncMock
    from app.services.oscal_service import ValidationResult, ValidationIssue, ConversionResult
    
    mock_service = Mock()
    
    # Mock validation
    mock_service.validate_document = AsyncMock(return_value=ValidationResult(
        is_valid=True,
        document_type="system-security-plan",
        errors=[],
        warnings=[],
        duration_ms=500,
        cli_stdout="Validation passed",
        cli_stderr="",
        return_code=0
    ))
    
    # Mock conversion
    mock_service.convert_document = AsyncMock(return_value=ConversionResult(
        success=True,
        output_path=Path("/tmp/converted.json"),
        input_format="xml",
        output_format="json",
        duration_ms=300,
        cli_stdout="Conversion successful",
        cli_stderr="",
        return_code=0
    ))
    
    mock_service.verify_cli_available = AsyncMock(return_value=True)
    
    return mock_service


@pytest.fixture
def mock_fedramp_service():
    """Mock FedRAMP service for testing."""
    from unittest.mock import Mock, AsyncMock
    from app.services.fedramp_service import FedRAMPValidationResult, FedRAMPValidationIssue
    
    mock_service = Mock()
    mock_service.validate_document = AsyncMock(return_value=FedRAMPValidationResult(
        is_compliant=True,
        baseline="moderate",
        document_type="system-security-plan",
        issues=[],
        validation_time_ms=1000,
        metadata={"controls_validated": 25}
    ))
    
    mock_service.get_baseline_requirements = AsyncMock(return_value={
        "required_controls": ["ac-1", "ac-2", "ac-3"],
        "min_controls": 3,
        "required_metadata": ["system_name", "system_id"]
    })
    
    return mock_service


class TestHelpers:
    """Test helper utilities."""
    
    @staticmethod
    def assert_operation_success(response_data: Dict[str, Any]) -> None:
        """Assert that an operation completed successfully."""
        assert "operation_id" in response_data
        assert response_data.get("success", False) is True
        
    @staticmethod
    def assert_validation_response(response_data: Dict[str, Any], should_be_valid: bool = True) -> None:
        """Assert validation response structure."""
        assert "is_valid" in response_data
        assert response_data["is_valid"] == should_be_valid
        assert "summary" in response_data
        assert "errors" in response_data
        
    @staticmethod
    def create_upload_file(content: str, filename: str = "test.json", content_type: str = "application/json"):
        """Create a file-like object for upload testing."""
        from io import StringIO
        return StringIO(content)


# Make TestHelpers available as a fixture
@pytest.fixture
def helpers():
    """Test helper utilities fixture."""
    return TestHelpers